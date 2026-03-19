import os
import shutil
import gc
from docx import Document
from docx.shared import Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf2docx import Converter

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_DIR = os.path.join(BASE_DIR, 'temp')
os.makedirs(TEMP_DIR, exist_ok=True)

def process():
    for file in os.listdir(BASE_DIR):
        file_path = os.path.join(BASE_DIR, file)
        if os.path.isdir(file_path) or file.lower().endswith('.py') or file == 'temp':
            continue
            
        ext = file.lower()
        name_no_ext = os.path.splitext(file)[0]
        output_file = os.path.join(BASE_DIR, f"{name_no_ext}.docx")

        try:
            success = False
            if ext.endswith((".jpg", ".jpeg", ".png", ".bmp", ".webp")):
                doc = Document()
                section = doc.sections[0]
                section.page_height = Mm(297)
                section.page_width = Mm(210)
                
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                
                run.add_picture(file_path, width=Inches(6.0))
                
                doc.save(output_file)
                success = True
                del doc
                
            elif ext.endswith(".pdf"):
                cv = Converter(file_path)
                cv.convert(output_file)
                cv.close()
                success = True
                del cv

            if success:
                shutil.move(file_path, os.path.join(TEMP_DIR, file))
                print(f"Processed: {file}")
            
            gc.collect()
        except Exception as e:
            print(f"Error {file}: {e}")

if __name__ == "__main__":
    process()
